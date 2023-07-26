from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

# Create a new document
doc = Document()

# Add a section to the document
section = doc.sections[0]

# Set the header
header = section.header

# Create a paragraph for the header
header_paragraph = header.paragraphs[0]

# Create a left-aligned run for the left header
left_run = header_paragraph.add_run()
left_run.text = "Left Header"
left_run.bold = True
left_run.paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Create a right-aligned run for the right header
right_run = header_paragraph.add_run()
right_run.text = "Right Header"
right_run.bold = True
right_run.paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Set the tab stops for the paragraph
tab_stops = header_paragraph.paragraph_format.tab_stops
tab_stops.add_tab_stop(Pt(432), WD_ALIGN_PARAGRAPH.RIGHT)  # Set the tab stop position for right alignment

# Add a tab character between the left and right headers
header_paragraph.add_run().text = "\t"

# Save the document
doc.save("document.docx")
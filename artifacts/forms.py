# -*- coding: utf-8 -*-

################################################################################
## Form generated from reading UI file 'form.ui'
##
## Created by: Qt User Interface Compiler version 6.5.0
##
## WARNING! All changes made in this file will be lost when recompiling UI file!
################################################################################
import os
import numpy as np
from datetime import datetime
import datetime as dt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import load_workbook


from database import Database
from configparser import ConfigParser
import pymysql.connections as MySQLdb
from PyQt6.uic import loadUi
from PyQt6.QtWidgets import QMessageBox, QDialog, QVBoxLayout, QFileDialog, QWidget, QTableWidgetItem
from PyQt6.QtCore import QUrl
from PyQt6.QtGui import QDesktopServices

from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from matplotlib.backends.backend_qt5agg import NavigationToolbar2QT as NavigationToolbar
import matplotlib.pyplot as plt

class insertform(QDialog):
    def __init__(self):
        """Initializes an instance of the insertfrom class.
        Usage:
        -----------
            c_insert = insertform()
        Return:
        -----------
            None
        """
        super(insertform, self).__init__()
        self.ui = loadUi(os.path.join(os.path.dirname(__file__), "Data_insertion.ui"), self)
        
        self.db = Database()
        config_object = ConfigParser()
        config_object.read("config.ini")
        userInfo = config_object["USERINFO"]
        self.LOGGED_USER_ID = userInfo["LOGGED_USER_ID"]
        self.LOGGED_USER_ROLE = userInfo["LOGGED_USER_ROLE"]
        self.ui.updateButton.setVisible(False)
        self.ui.insertButton.clicked.connect(self.insertData)
        self.ui.cancelButton.clicked.connect(self.cancel)
        self.loadBox()

    def onBoxTextChanged(self, text):
        """
        """
        conn = MySQLdb.Connection(host=self.db.DB_SERVER, user=self.db.DB_USERNAME, password=self.db.DB_PASSWORD,
                                   database=self.db.DB_NAME)
        cursor = conn.cursor()
        query_role_select = f"SELECT role FROM managers WHERE id='{self.LOGGED_USER_ID}'"
        cursor.execute(query_role_select)
        data_role = cursor.fetchone()
        if data_role[0] =="user" and text == "Consumable":
            self.actionBox.setEnabled(False)
        elif data_role[0] =="user" and text == "Tooling":
            self.actionBox.setEnabled(True)

    def loadBox(self):
        """
        """
        self.ui.TypeBox.currentTextChanged.connect(self.onBoxTextChanged)
        # Add list of project in comboBox
        conn = MySQLdb.Connection(host=self.db.DB_SERVER, user=self.db.DB_USERNAME, password=self.db.DB_PASSWORD,
                                   database=self.db.DB_NAME)
        cursor = conn.cursor()
        query_project_select = f"SELECT projectname FROM projects p WHERE p.user_id='{self.LOGGED_USER_ID}'"
        query_action_select = f"SELECT Distinct actions FROM object_dist"
        query_category_select = f"SELECT Distinct Category FROM object_dist"
        query_role_select = f"SELECT role FROM managers WHERE id='{self.LOGGED_USER_ID}'"
        cursor.execute(query_project_select)
        data_proj = cursor.fetchall()
        cursor4 =conn.cursor()
        cursor4.execute(query_role_select)
        data_role = cursor4.fetchone()
        
        for i in data_proj : 
            res = i[0]
            self.projectBox.addItem(res)

        cursor2 = conn.cursor()
        cursor2.execute(query_action_select)
        data_action = cursor2.fetchall()
        for i in data_action : 
            res = i[0]
            self.actionBox.addItem(res)

        cursor5 = conn.cursor()
        query_type_select = f"SELECT distinct Type_object FROM object_dist"
        cursor5.execute(query_type_select)
        data_type = cursor.fetchall()
        for i in data_type : 
            res = i[0]
            self.ui.TypeBox.addItem(res)

        if data_role[0] == "user" and self.ui.TypeBox.currentText() == "Consumable":
            self.actionBox.setEnabled(False)

        cursor3 = conn.cursor()
        cursor3.execute(query_category_select)
        data_cat = cursor3.fetchall()
        for i in data_cat : 
            res = i[0]
            self.categoryBox.addItem(res)

    def cancel(self):
        """This function allows the user to navigate back to the previous screen in the application.
        Usage:
        -----------
            ui = Ui_ForgetPassword()
            ui.goback()

        Return:
        -----------
            None
        """
        self.close()

    def insertData(self):
        """
        This function will insert a row of data in the object_dist table
        :return:
        """
        objectText      = str(self.objectText.text())
        typeText        = str(self.TypeBox.currentText())
        locationText    = str(self.locationText.text())
        calibrationText = str(self.calibrationText.text())
        quantityText    = str(self.quantityText.text())
        actionText      = str(self.actionBox.currentText())
        projectText     = str(self.projectBox.currentText())
        categoryText    = str(self.categoryBox.currentText())

        if(len(objectText.strip()) == 0 or len(calibrationText.strip()) == 0 or len(quantityText.strip()) == 0 or len(locationText.strip()) == 0 or len(typeText.strip()) == 0
            or len(projectText.strip()) == 0):
            msgBox = QMessageBox()
            msgBox.setText("All fields must be filled before confirming !")
            msgBox.setWindowTitle("Attention")
            msgBox.exec()
        else:
            input_insert = (objectText, typeText, locationText, calibrationText, projectText, categoryText)
            conn = MySQLdb.Connection(host=self.db.DB_SERVER, user=self.db.DB_USERNAME, password=self.db.DB_PASSWORD,
                                database=self.db.DB_NAME)
            mycursor = conn.cursor()
            query = "SELECT Object, Type_Object, location, CAST(Calibration as char), project_name, Category FROM object_dist"
            mycursor.execute(query)
            data = mycursor.fetchall()
            if input_insert in data and actionText == "Adding":
                    #UPDATE INST TABLE
                    query_update = "UPDATE object_dist SET Quantity= Quantity + %s WHERE Object = '%s' AND Type_Object='%s' AND Location ='%s' AND Calibration =%s AND project_name ='%s' "  % (int(quantityText), objectText, typeText, locationText, int(calibrationText), projectText)
                    mycursor.execute(query_update)
                    conn.commit()
                    self.close()
                    
                    #INSERT INTO HISTORIQUE TABLE
                    now = datetime.now()
                    date_time = now.strftime("%Y-%m-%d %H:%M:%S")
                    query_insert = "INSERT INTO historique(Object,Type_object,Location, Calibration, Quantity, Category,operation_datetime,user_id,actions,project_name) VALUES ('%s','%s','%s','%s','%s','%s','%s','%s','%s','%s')" %(objectText,typeText,locationText,calibrationText,quantityText,categoryText, date_time, self.LOGGED_USER_ID, actionText, projectText)
                    mycursor.execute(query_insert)
                    conn.commit()
                    self.close()
            elif input_insert in data and actionText == "Take Out":
                    query_update = "UPDATE object_dist SET Quantity= Quantity - %s WHERE Object = '%s' AND Type_Object='%s' AND Location ='%s' AND Calibration =%s AND project_name ='%s' "  % (int(quantityText), objectText, typeText, locationText, int(calibrationText), projectText)
                    mycursor.execute(query_update)
                    query_delete = "DELETE FROM object_dist WHERE Quantity =0"
                    mycursor.execute(query_delete)
                    conn.commit()
                    self.close() 

                    #INSERT INTO HISTORIQUE TABLE
                    now = datetime.now()
                    date_time = now.strftime("%Y-%m-%d %H:%M:%S")
                    query_insert = "INSERT INTO historique(Object,Type_object,Location, Calibration, Quantity, Category,operation_datetime,user_id,actions,project_name) VALUES ('%s','%s','%s','%s','%s','%s','%s','%s','%s','%s')" %(objectText,typeText,locationText,calibrationText,quantityText,categoryText, date_time, self.LOGGED_USER_ID, actionText, projectText)
                    mycursor.execute(query_insert)
                    conn.commit()
                    self.close()
            else:
                    now = datetime.now()
                    date_time = now.strftime("%Y-%m-%d %H:%M:%S")
                    conn = MySQLdb.Connection(host=self.db.DB_SERVER, user=self.db.DB_USERNAME, password=self.db.DB_PASSWORD,
                                                        database=self.db.DB_NAME)
                    mycursor = conn.cursor()
                    query = "SELECT Object, Type_Object, location, Calibration, Quantity FROM object_dist"
                    mycursor.execute(query)
                    data = mycursor.fetchall()

                    #INSERT INTO INST TABLE
                    query_insert = "INSERT INTO object_dist(Object,Type_object,Location, Calibration, Quantity, Category,operation_datetime,user_id,actions,project_name) VALUES ('%s','%s','%s','%s','%s','%s','%s','%s','%s','%s')" %(objectText,typeText,locationText,calibrationText,quantityText,categoryText, date_time, self.LOGGED_USER_ID, actionText, projectText)
                    mycursor.execute(query_insert)
                    conn.commit()
                    self.close()

                    #INSERT INTO HISTORIQUE TABLE
                    query_insert_hist = "INSERT INTO historique(Object,Type_object,Location, Calibration, Quantity, Category,operation_datetime,user_id,actions,project_name) VALUES ('%s','%s','%s','%s','%s','%s','%s','%s','%s','%s')" %(objectText,typeText,locationText,calibrationText,quantityText,categoryText, date_time, self.LOGGED_USER_ID, actionText, projectText)
                    mycursor.execute(query_insert_hist)
                    conn.commit()
                    self.close()
                      
class profileform(QDialog):
    def __init__(self):
        """Initializes an instance of the Ui_SignUp class.
        Usage:
        -----------
            ui = Ui_SignUp()
        Return:
        -----------
            None
        """
        super(profileform, self).__init__()
        loadUi(os.path.join(os.path.dirname(__file__), "profile.ui"), self)

        self.saveButton.clicked.connect(self.updateData)

        self.db = Database()
        config_object = ConfigParser()
        config_object.read("config.ini")
        userInfo = config_object["USERINFO"]
        self.LOGGED_USER_ID = userInfo["LOGGED_USER_ID"]
        try:
            conn = MySQLdb.Connection(host=self.db.DB_SERVER, user=self.db.DB_USERNAME, password=self.db.DB_PASSWORD,
                                   database=self.db.DB_NAME)
            cursor = conn.cursor()
            query = f"SELECT username,email,firstname,lastname,ismanager, Society, Entity FROM managers WHERE id='{self.LOGGED_USER_ID}'"
            query_proj = f"SELECT projectname FROM projects p INNER JOIN managers m ON p.user_id = m.id  WHERE m.id='{self.LOGGED_USER_ID}'"
            cursor.execute(query)
            data = cursor.fetchone()
            cursor.execute(query_proj)
            data_proj = cursor.fetchall()
            self.usernameTextField.setText(data[0])
            self.emailTextField.setText(data[1])
            self.firstNameTextField.setText(data[2])
            self.lastNameTextField.setText(data[3])
            self.usernameTextField_2.setText(data[5])
            self.usernameTextField_3.setText(data[6])
            for i in data_proj : 
                res = i[0]
                self.comboBox.addItem(res)
            if(int(data[4]) == 0):
                self.usernameTextField.setEnabled(False)
                self.emailTextField.setEnabled(False)
                self.firstNameTextField.setEnabled(False)
                self.lastNameTextField.setEnabled(False)
                self.saveButton.setEnabled(False)
                self.usernameTextField_2.setEnabled(False)
                self.usernameTextField_3.setEnabled(False)
        except Exception as e:
            QMessageBox.about(str(e))

    def updateData(self):
        """
        """
        username = str(self.usernameTextField.text())
        email    = str(self.emailTextField.text())
        firstName= str(self.firstNameTextField.text())
        lastName = str(self.lastNameTextField.text())

        conn = MySQLdb.Connection(host=self.db.DB_SERVER, user=self.db.DB_USERNAME, password=self.db.DB_PASSWORD,
                               database=self.db.DB_NAME)
        mycursor = conn.cursor()
        query = "UPDATE managers SET username='%s',email='%s',firstname='%s',lastname='%s' WHERE id=%s" % (
        username,email,firstName,lastName,self.LOGGED_USER_ID)
        mycursor.execute(query)
        conn.commit()
        self.close()

class projectform(QDialog):
    def __init__(self):
        """Initializes an instance of the Ui_SignUp class.
        Usage:
        -----------
            ui = Ui_SignUp()
        Return:
        -----------
            None
        """
        super(projectform, self).__init__()
        self.ui = loadUi(os.path.join(os.path.dirname(__file__), "proj_form.ui"), self)

        self.figure = plt.figure()
        self.canvas = FigureCanvas(self.figure)
        self.toolbar = NavigationToolbar(self.canvas, self)
        layout = QVBoxLayout()
        layout.addWidget(self.toolbar)
        layout.addWidget(self.canvas)
        self.setLayout(layout)

        self.db = Database()
        config_object = ConfigParser()
        config_object.read("config.ini")
        userInfo = config_object["USERINFO"]
        self.LOGGED_USER_ID = userInfo["LOGGED_USER_ID"]
   
        self.plotActivity()
    
    def convertTime(self, date_string):
        """
        """
        date_format = '%Y-%m-%d %H:%M:%S'
        date = dt.datetime.strptime(date_string, date_format)
        week_number = date.strftime('%U')
        year = date.strftime('%Y')
        week_year_format = f'S{week_number}-{year}'
        return week_year_format
    
    def calculateSumByDate(self,date, act):
        result = {}
        for d, a in zip(date, act):
            if d in result:
                result[d] += a
            else:
                result[d] = a
        return result
    
    def addMissingElements(self, weeks_list, weeks_target, act_target):
        for week in weeks_list:
            if week not in weeks_target:
                index = weeks_list.index(week)
                weeks_target.insert(index, week)
                act_target.insert(index, 0)
    
    def plotActivity(self):
        """
        """
        #-------------------------------------------Adding activity----------------------------------------------#
        conn = MySQLdb.Connection(host=self.db.DB_SERVER, user=self.db.DB_USERNAME, password=self.db.DB_PASSWORD,
                                   database=self.db.DB_NAME)
        cursor = conn.cursor()
        query_add = f"SELECT operation_datetime, Quantity FROM object_dist WHERE actions = 'Adding' AND user_id='{self.LOGGED_USER_ID}' Order by operation_datetime"
        cursor.execute(query_add)
        table_add = cursor.fetchall()
        ss_aaaa_add = []
        activ_add =[]
        for i in range(len(table_add)):
            ss_aaaa_add.append(self.convertTime(str(table_add[i][0])))
            activ_add.append(table_add[i][1])
        
        result = self.calculateSumByDate(ss_aaaa_add, activ_add)
        weeks = []
        activities =[]
        for i in result.keys():
            weeks.append(i)
            activities.append(result[i])
        
        #-------------------------------------------Take out activity--------------------------------------------#
        query_take = f"SELECT operation_datetime, Quantity FROM object_dist WHERE actions = 'Take Out' AND user_id='{self.LOGGED_USER_ID}' Order BY operation_datetime"
        cursor.execute(query_take)
        table_take = cursor.fetchall()
        ss_aaaa_take = []
        activ_take =[]
        for i in range(len(table_take)):
            ss_aaaa_take.append(self.convertTime(str(table_take[i][0])))
            activ_take.append(table_take[i][1])
        
        result_take = self.calculateSumByDate(ss_aaaa_take, activ_take)
        weeks_take = []
        activities_take =[]
        for i in result_take.keys():
            weeks_take.append(i)
            activities_take.append(result_take[i])

        self.addMissingElements(weeks, weeks_take, activities_take)
        self.addMissingElements(weeks_take, weeks, activities)
        self.addMissingElements(ss_aaaa_add, ss_aaaa_take, activ_take)
        self.addMissingElements(ss_aaaa_take, ss_aaaa_add, activ_add)

        unique_weeks = np.unique(weeks)
        unique_ss_aaaa_add = np.unique(ss_aaaa_add)
        #-----------------------------------------------------Plotting-------------------------------------------------#    
        self.figure.clear()
        ax = self.figure.add_subplot(121)
        ax.plot(weeks, activities, 'o-', label='Adding action')
        ax.plot(weeks_take, activities_take, 'o-', label='Take out action')
        ax.set_xticklabels(unique_weeks , rotation=45)
        
        ax.set_title('Activity by SS-AAAA')
        ax.set_ylabel('Quantity')
        ax.legend()

        ax = self.figure.add_subplot(122)
        ax.scatter(ss_aaaa_add, activ_add, label='Adding action')
        ax.scatter(ss_aaaa_take, activ_take, label='Take out action')
        ax.set_title('Activity by SS-AAAA')
        ax.set_ylabel('Quantity')
        ax.set_xticklabels(unique_ss_aaaa_add , rotation=45)
        ax.legend() 
        self.canvas.draw()

class tendency(QDialog):
    def __init__(self):
        """
        """
        super(tendency, self).__init__()
        self.ui = loadUi(os.path.join(os.path.dirname(__file__), "plot_activ.ui"), self)

        self.db = Database()
        config_object = ConfigParser()
        config_object.read("config.ini")
        userInfo = config_object["USERINFO"]
        self.LOGGED_USER_ID = userInfo["LOGGED_USER_ID"]

        self.figure = plt.figure()
        self.canvas = FigureCanvas(self.figure)
        self.toolbar = NavigationToolbar(self.canvas, self)
        layout = QVBoxLayout(self.ui.widget)
        layout.addWidget(self.toolbar)
        layout.addWidget(self.canvas)

        self.figure2 = plt.figure()
        self.canvas2 = FigureCanvas(self.figure2)
        self.toolbar2 = NavigationToolbar(self.canvas2, self)
        self.layout2 = QVBoxLayout(self.ui.widget_2)
        self.layout2.addWidget(self.toolbar2)
        self.layout2.addWidget(self.canvas2)

        
        self.figure3 = plt.figure()
        self.canvas3 = FigureCanvas(self.figure3)
        self.toolbar3 = NavigationToolbar(self.canvas3, self)
        self.layout3 = QVBoxLayout(self.ui.widget_3)
        self.layout3.addWidget(self.toolbar3)
        self.layout3.addWidget(self.canvas3)

        self.figure4 = plt.figure()
        self.canvas4 = FigureCanvas(self.figure4)
        self.toolbar4 = NavigationToolbar(self.canvas4, self)
        self.layout4 = QVBoxLayout(self.ui.widget_4)
        self.layout4.addWidget(self.toolbar4)
        self.layout4.addWidget(self.canvas4)
        

        self.loadBox(self.ui.project_box)
        self.loadBox(self.ui.person_box, type="person")
        self.loadBox(self.ui.object_box, type="object")
        self.projectsHist()
        self.plotUsers()
        self.plotObject()
        self.plotQuantity()

        self.ui.project_box.currentIndexChanged.connect(self.highlightProjectsHist)
        self.ui.person_box.currentIndexChanged.connect(self.highlightPersonHist)
        self.ui.object_box.currentIndexChanged.connect(self.highlightObjectHist)
        self.ui.export_btn.clicked.connect(self.exportWordDoc_2)

    def loadBox(self, box, type="project"):
        conn = MySQLdb.Connection(host=self.db.DB_SERVER, user=self.db.DB_USERNAME, password=self.db.DB_PASSWORD,
                                   database=self.db.DB_NAME)
        cursor = conn.cursor()
        if type=="project":
            query = f"SELECT Distinct projectname FROM projects"
        elif type=="person":
            query = f"SELECT Distinct m.name FROM managers m"
        elif type =="object":
            query =  f"SELECT Distinct Object FROM object_dist"
        cursor.execute(query)
        data = cursor.fetchall()
        box.addItem("All")
        for i in data:
            res = i[0]
            box.addItem(res)

    def projectsHist(self, condition="1=1"):
        conn = MySQLdb.Connection(host=self.db.DB_SERVER, user=self.db.DB_USERNAME, password=self.db.DB_PASSWORD,
                                  database=self.db.DB_NAME)
        cursor = conn.cursor()
        query = f"SELECT sum(cout), project_name FROM fakegcino.object_dist WHERE {condition} GROUP BY project_name "
        cursor.execute(query)
        table_data = cursor.fetchall()
        amounts = [float(row[0]) for row in table_data if isinstance(row[0], (int, float))]
        projects = [row[1] for row in table_data]

        # Créer un graphique à barres en utilisant Pyqtgraph
        if (condition == "1=1"):
            self.ui.widget_2.setVisible(False)
            self.ui.widget_3.setVisible(False)
            self.ui.widget_4.setVisible(False)
            self.ui.person_box.setEnabled(False)
            self.ui.object_box.setEnabled(False)
            self.ui.export_btn.setEnabled(True)
            self.figure.clear()
            ax1 = self.figure.add_subplot(111)
            bar1 = ax1.bar(projects, amounts, width=0.35, align='center', alpha=0.4, color='b')
            for rect in bar1:
                height = rect.get_height()
                ax1.text(rect.get_x() + rect.get_width() / 2.0, height, f'{height:.0f}'+'€', ha='center', va='top')
            ax1.set_title('Amounts by projects')
            ax1.set_ylabel('Amounts')
            self.canvas.draw()
        elif (condition == "2=2"):
            self.ui.widget_2.setVisible(False)
            self.ui.widget_3.setVisible(False)
            self.ui.widget_4.setVisible(False)
            self.ui.person_box.setEnabled(False)
            self.ui.object_box.setEnabled(False)
            self.ui.export_btn.setEnabled(True)
            self.figure.clear()
            ax1 = self.figure.add_subplot(111)
            bar1 = ax1.bar(projects, amounts, width=0.35, align='center', alpha=0.4, color='b')
            for rect in bar1:
                height = rect.get_height()
                ax1.text(rect.get_x() + rect.get_width() / 2.0, height, f'{height:.0f}'+'€', ha='center', va='top')
            ax1.set_title('Amounts by projects')
            ax1.set_ylabel('Amounts')
            self.canvas.draw()
        else:
            self.ui.widget_2.setVisible(True)
            self.ui.widget_3.setVisible(True)
            self.ui.widget_4.setVisible(True)
            self.ui.person_box.setEnabled(True)
            self.ui.object_box.setEnabled(True)
            self.ui.export_btn.setEnabled(True)
            self.figure.clear()
            ax1 = self.figure.add_subplot(111)
            bar1 =ax1.bar(projects, amounts, width=0.35, align='center', alpha=0.4, color='b')
            for rect in bar1:
                height = rect.get_height()
                ax1.text(rect.get_x() + rect.get_width() / 2.0, height, f'{height:.0f}'+'€', ha='center', va='top')
            ax1.set_title('Amounts by projects')
            ax1.set_ylabel('Amounts')
            self.canvas.draw()
        
        # Save the bar chart as an image file
        self.graph_filename1 = os.path.dirname(os.path.dirname(__file__))+ "\\run\images\projects.png"
        self.figure.savefig(self.graph_filename1)
        
    def plotUsers(self, condition="1=1"):
        """
        """
        conn = MySQLdb.Connection(host=self.db.DB_SERVER, user=self.db.DB_USERNAME, password=self.db.DB_PASSWORD,
                                  database=self.db.DB_NAME)
        cursor = conn.cursor()
        query = f"SELECT sum(o.cout), m.name FROM fakegcino.object_dist o INNER JOIN fakegcino.managers m ON o.user_id = m.id WHERE o.project_name ='{self.ui.project_box.currentText()}' GROUP BY m.name "
        cursor.execute(query)
        table_data = cursor.fetchall()
        amounts = [float(row[0]) for row in table_data if isinstance(row[0], (int, float))]
        projects = [row[1] for row in table_data]

        if condition =="1=1":
            ax = self.figure2.add_subplot(111)
            bar1 = ax.bar(projects, amounts, width=0.35, align='center', alpha=0.4, color='r')
            for rect in bar1:
                height = rect.get_height()
                ax.text(rect.get_x() + rect.get_width() / 2.0, height, f'{height:.0f}'+'€', ha='center', va='top')
            ax.set_title('Amounts by users')
            ax.set_ylabel('Amounts')
            self.canvas2.draw()

        elif (condition == "2=2"): 
            self.figure2.clear()
            ax = self.figure2.add_subplot(111)
            bar1 = ax.bar(projects, amounts, width=0.35, align='center', alpha=0.4, color='r')
            for rect in bar1:
                height = rect.get_height()
                ax.text(rect.get_x() + rect.get_width() / 2.0, height, f'{height:.0f}'+'€', ha='center', va='top')
            ax.set_title('Amounts by users')
            ax.set_ylabel('Amounts')
            self.canvas2.draw()
           
        else:
             if self.ui.person_box.currentText() != 'All' :
                query = f"SELECT sum(o.cout), m.name FROM fakegcino.object_dist o INNER JOIN fakegcino.managers m ON o.user_id = m.id WHERE o.project_name ='{self.ui.project_box.currentText()}' AND m.name = '{self.ui.person_box.currentText()}' GROUP BY m.name "       
             else:
                 query = query   
                      
             cursor.execute(query)
             table_data = cursor.fetchall()
             amounts = [float(row[0]) for row in table_data if isinstance(row[0], (int, float))]
             projects = [row[1] for row in table_data]
             self.figure2.clear()
             ax = self.figure2.add_subplot(111)
             bar1 = ax.bar(projects, amounts, width=0.35, align='center', alpha=0.4, color='r')
             for rect in bar1:
                height = rect.get_height()
                ax.text(rect.get_x() + rect.get_width() / 2.0, height, f'{height:.0f}'+'€', ha='center', va='top')
             ax.set_title('Amounts by users')
             ax.set_ylabel('Amounts')
             self.canvas2.draw()
        # Save the bar chart as an image file
        self.graph_filename2 = os.path.dirname(os.path.dirname(__file__))+ r"\run\images\users.png"
        self.figure2.savefig(self.graph_filename2)

    def plotObject(self, condition ="1=1"):
        conn = MySQLdb.Connection(host=self.db.DB_SERVER, user=self.db.DB_USERNAME, password=self.db.DB_PASSWORD,
                                  database=self.db.DB_NAME)
        cursor = conn.cursor()
        if self.ui.person_box.currentText() == 'All':
            query = f"SELECT sum(o.cout), o.Object FROM fakegcino.object_dist o INNER JOIN fakegcino.managers m ON m.id = o.user_id Where o.project_name ='{self.ui.project_box.currentText()}' Group BY Object"
        else:
            query = f"SELECT sum(o.cout), o.Object FROM fakegcino.object_dist o INNER JOIN fakegcino.managers m ON m.id = o.user_id Where o.project_name ='{self.ui.project_box.currentText()}' AND m.name = '{self.ui.person_box.currentText()}' Group BY Object"
        cursor.execute(query)
        table_data = cursor.fetchall()
        amounts = [float(row[0]) for row in table_data if isinstance(row[0], (int, float))]
        projects = [row[1] for row in table_data]

        if condition =="1=1":
            ax = self.figure3.add_subplot(111)
            bar1 = ax.bar(projects, amounts, width=0.75, align='center', alpha=0.4, color='g')
            for rect in bar1:
                height = rect.get_height()
                ax.text(rect.get_x() + rect.get_width() / 2.0, height, f'{height:.0f}'+"€", ha='center', va='top')
            ax.set_title('Amounts by objects')
            ax.set_ylabel('Amounts')
            self.figure3.autofmt_xdate(rotation=45)
            self.canvas3.draw()
        elif (condition == "2=2"): 
            self.figure3.clear()
            ax = self.figure3.add_subplot(111)
            bar1 = ax.bar(projects, amounts, width=0.75, align='center', alpha=0.4, color='g')
            for rect in bar1:
                height = rect.get_height()
                ax.text(rect.get_x() + rect.get_width() / 2.0, height, f'{height:.0f}'+"€", ha='center', va='top')
            ax.set_title('Amounts by objects')
            ax.set_ylabel('Amounts')
           
            for rect in bar1:
                height = rect.get_height()
                ax.text(rect.get_x() + rect.get_width() / 2.0, height, f'{height:.0f}'+"€", ha='center', va='top')
            self.figure3.autofmt_xdate(rotation=45)
            self.canvas3.draw()
        else: 
            if self.ui.object_box.currentText() != 'All' and self.ui.person_box.currentText() !='All' :
                query = f"SELECT sum(cout), Object FROM fakegcino.object_dist o INNER JOIN fakegcino.managers m ON m.id = o.user_id Where project_name ='{self.ui.project_box.currentText()}' AND m.name = '{self.ui.person_box.currentText()}'AND object='{self.ui.object_box.currentText()}' Group BY Object "       
            elif self.ui.object_box.currentText() !='All' and self.ui.person_box.currentText() =='All':
                query = f"SELECT sum(cout), Object FROM fakegcino.object_dist o INNER JOIN fakegcino.managers m ON m.id = o.user_id Where project_name ='{self.ui.project_box.currentText()}' AND object='{self.ui.object_box.currentText()}' Group BY Object "
            else:
                 query = query 
            cursor.execute(query)
          
            table_data = cursor.fetchall()
            amounts = [float(row[0]) for row in table_data if isinstance(row[0], (int, float))]
            projects = [row[1] for row in table_data]

            self.figure3.clear()
            ax = self.figure3.add_subplot(111)
            bar1 = ax.bar(projects, amounts, width=0.75, align='center', alpha=0.4, color='g')
            for rect in bar1:
                height = rect.get_height()
                ax.text(rect.get_x() + rect.get_width() / 2.0, height, f'{height:.0f}'+"€", ha='center', va='top')
            ax.set_title('Amounts by objects')
            ax.set_ylabel('Amounts')
            self.figure3.autofmt_xdate(rotation=45)
            self.canvas3.draw()
        # Save the bar chart as an image file
        self.graph_filename3 = os.path.dirname(os.path.dirname(__file__))+ "\\run\images\objects.png"
        self.figure3.savefig(self.graph_filename3)

    def addMissingElements(self, weeks_list, weeks_target, act_target):
        for week in weeks_list:
            if week not in weeks_target:
                index = weeks_list.index(week)
                weeks_target.insert(index, week)
                act_target.insert(index, 0)

    def plotQuantity(self, condition="1=1"):
        """
        """
        #-------------------------------------------Adding activity----------------------------------------------#
        conn = MySQLdb.Connection(host=self.db.DB_SERVER, user=self.db.DB_USERNAME, password=self.db.DB_PASSWORD,
                                   database=self.db.DB_NAME)
        cursor = conn.cursor()
        if self.ui.person_box.currentText() == 'All':
            if self.ui.object_box.currentText() == 'All':
                query_add = f"SELECT operation_datetime, Quantity FROM object_dist o INNER JOIN fakegcino.managers m ON m.id = o.user_id Where o.project_name ='{self.ui.project_box.currentText()}' AND actions = 'Adding' Group BY operation_datetime, Quantity order by operation_datetime"
                query_take = f"SELECT operation_datetime, Quantity FROM object_dist o INNER JOIN fakegcino.managers m ON m.id = o.user_id Where o.project_name ='{self.ui.project_box.currentText()}' AND actions = 'Take Out' Group BY operation_datetime, Quantity order by operation_datetime"       
            else:
                query_add = f"SELECT operation_datetime, Quantity FROM object_dist o INNER JOIN fakegcino.managers m ON m.id = o.user_id Where o.project_name ='{self.ui.project_box.currentText()}'  AND o.Object = '{self.ui.object_box.currentText()}' AND actions = 'Take Out' Group BY operation_datetime, Quantity order by operation_datetime"
                query_take = f"SELECT operation_datetime, Quantity FROM object_dist o INNER JOIN fakegcino.managers m ON m.id = o.user_id Where o.project_name ='{self.ui.project_box.currentText()}' AND o.Object = '{self.ui.object_box.currentText()}' AND actions = 'Adding' Group BY operation_datetime, Quantity order by operation_datetime"
        else: 
            if self.ui.object_box.currentText() == 'All':
                query_add = f"SELECT operation_datetime, Quantity FROM object_dist o INNER JOIN fakegcino.managers m ON m.id = o.user_id Where o.project_name ='{self.ui.project_box.currentText()}'  AND m.name = '{self.ui.person_box.currentText()}' AND actions = 'Adding' Group BY operation_datetime, Quantity order by operation_datetime"
                query_take = f"SELECT operation_datetime, Quantity FROM object_dist o INNER JOIN fakegcino.managers m ON m.id = o.user_id Where o.project_name ='{self.ui.project_box.currentText()}'  AND m.name = '{self.ui.person_box.currentText()}' AND actions = 'Take Out' Group BY operation_datetime, Quantity order by operation_datetime"       
            else:
                query_add = f"SELECT operation_datetime, Quantity FROM object_dist o INNER JOIN fakegcino.managers m ON m.id = o.user_id Where o.project_name ='{self.ui.project_box.currentText()}' AND m.name = '{self.ui.person_box.currentText()}' AND o.Object = '{self.ui.object_box.currentText()}' AND actions = 'Take Out' Group BY operation_datetime, Quantity order by operation_datetime"
                query_take = f"SELECT operation_datetime, Quantity FROM object_dist o INNER JOIN fakegcino.managers m ON m.id = o.user_id Where o.project_name ='{self.ui.project_box.currentText()}' AND m.name = '{self.ui.person_box.currentText()}' AND o.Object = '{self.ui.object_box.currentText()}' AND actions = 'Adding' Group BY operation_datetime, Quantity order by operation_datetime"
        
        cursor.execute(query_add)
        table_add = cursor.fetchall()
        ss_aaaa_add = []
        activ_add =[]
        for i in range(len(table_add)):
            ss_aaaa_add.append(self.convertTime(str(table_add[i][0])))
            activ_add.append(table_add[i][1])
        
        result = self.calculateSumByDate(ss_aaaa_add, activ_add)
        weeks = []
        activities =[]
        for i in result.keys():
            weeks.append(i)
            activities.append(result[i])
        #-------------------------------------------Take out activity--------------------------------------------#
        #query_take = f"SELECT operation_datetime, Quantity FROM object_dist WHERE actions = 'Take Out' order by operation_datetime"
        cursor.execute(query_take)
        table_take = cursor.fetchall()
        ss_aaaa_take = []
        activ_take =[]
        for i in range(len(table_take)):
            ss_aaaa_take.append(self.convertTime(str(table_take[i][0])))
            activ_take.append(table_take[i][1])
        
        result_take = self.calculateSumByDate(ss_aaaa_take, activ_take)
        weeks_take = []
        activities_take =[]
        for i in result_take.keys():
            weeks_take.append(i)
            activities_take.append(result_take[i])
        self.addMissingElements(weeks, weeks_take, activities_take)
        self.addMissingElements(weeks_take, weeks, activities)
        unique_weeks = np.unique(weeks)
        ###################################Plotting
        if (condition == "1=1"):
            self.figure4.clear()
            ax = self.figure4.add_subplot(111)
            ax.plot(weeks, activities, 'o-', label='Adding action', color='#08F7FE')
            ax.plot(weeks_take, activities_take, 'o-', label='take out action', color='#FE53BB')
            ax.set_title('Quantity by Date')
            ax.set_ylabel('Quantity')
            ax.legend()
            ax.set_facecolor('#212946')
            ax.fill_between(weeks, activities,
                    color='#08F7FE',
                    alpha=0.1)
            ax.fill_between(weeks_take, activities_take,
                    color='#FE53BB',
                    alpha=0.1)
            ax.grid(color='#2A3459')
            for i in range(len(activities)):
                ax.text(weeks[i], activities[i], f"{activities[i]}",
                horizontalalignment='center',
                verticalalignment='bottom',
                color = '#08F7FE')
            for i in range(len(activities_take)):
                ax.text(weeks_take[i], activities_take[i], f"{activities_take[i]}",
                horizontalalignment='center',
                verticalalignment='top',
                color = '#FE53BB')
            ax.set_xticklabels(unique_weeks , rotation=20)
            self.canvas4.draw()

        elif (condition == "2=2"):
            self.figure4.clear()
            ax = self.figure4.add_subplot(111)
            ax.plot(weeks, activities, 'o-', label='Adding action', color='#08F7FE')
            ax.plot(weeks_take, activities_take, 'o-', label='take out action', color='#FE53BB')
            ax.set_title('Quantity by Date')
            ax.set_ylabel('Quantity')
            ax.legend()
            ax.set_facecolor('#212946')
            ax.fill_between(weeks, activities,
                    color='#08F7FE',
                    alpha=0.1)
            ax.fill_between(weeks_take, activities_take,
                    color='#FE53BB',
                    alpha=0.1)
            ax.grid(color='#2A3459')
            for i in range(len(activities)):
                ax.text(weeks[i], activities[i], f"{activities[i]}",
                horizontalalignment='center',
                verticalalignment='bottom',
                color = '#08F7FE')

            for i in range(len(activities_take)):
                ax.text(weeks_take[i], activities_take[i], f"{activities_take[i]}",
                horizontalalignment='center',
                verticalalignment='top',
                color = '#FE53BB')
            ax.set_xticklabels(unique_weeks , rotation=20)
            self.canvas4.draw()
        else:
            self.figure4.clear()
            ax = self.figure4.add_subplot(111)
            ax.plot(weeks, activities, 'o-', label='Adding action', color='#08F7FE')
            ax.plot(weeks_take, activities_take, 'o-', label='take out action', color='#FE53BB')
            ax.set_title('Quantity by Date')
            ax.set_ylabel('Quantity')
            ax.legend()
            ax.set_facecolor('#212946')
            ax.fill_between(weeks, activities,
                    color='#08F7FE',
                    alpha=0.1)
            ax.fill_between(weeks_take, activities_take,
                    color='#FE53BB',
                    alpha=0.1)
            ax.grid(color='#2A3459')
            for i in range(len(activities)):
                ax.text(weeks[i], activities[i], f"{activities[i]}",
                horizontalalignment='center',
                verticalalignment='bottom',
                color = '#08F7FE')

            for i in range(len(activities_take)):
                ax.text(weeks_take[i], activities_take[i], f"{activities_take[i]}",
                horizontalalignment='center',
                verticalalignment='top',
                color = '#FE53BB')
            ax.set_xticklabels(unique_weeks , rotation=10)
            self.canvas4.draw()
        # Save the bar chart as an image file
        self.graph_filename4 = os.path.dirname(os.path.dirname(__file__))+ "\\run\images\quantity.png"
        self.figure4.savefig(self.graph_filename4)

    def highlightProjectsHist(self):
        """
        """
        if self.ui.project_box.currentText()== "All":
            return self.projectsHist(condition="2=2"), self.plotUsers(condition = "2=2"), self.plotObject(condition = "2=2") , self.plotQuantity(condition = "2=2")
        else:
            return self.projectsHist(condition=f"Project_name = '{self.ui.project_box.currentText()}'"), self.plotUsers(condition="3=3"), self.plotObject(condition="3=3")  , self.plotQuantity(condition = "3=3")

    def highlightPersonHist(self):
        """
        """
        if self.ui.person_box.currentText()== "All":
            self.plotUsers(condition = "2=2"), self.plotObject(condition = "2=2"), self.plotQuantity(condition = "2=2")
        else:
            return self.plotUsers(condition="3=3"), self.plotObject(condition="3=3") , self.plotQuantity(condition = "3=3")
        
    def highlightObjectHist(self):
        """
        """
        if self.ui.object_box.currentText() == "All":
            self.plotObject(condition = "2=2"), self.plotQuantity(condition = "2=2")
        else:
            return self.plotObject(condition="3=3"), self.plotQuantity(condition = "3=3")

    def convertTime(self, date_string):
        """
        """
        date_format = '%Y-%m-%d %H:%M:%S'
        date = dt.datetime.strptime(date_string, date_format)
        week_number = date.strftime('%U')
        year = date.strftime('%Y')
        week_year_format = f'S{week_number}-{year}'
        return week_year_format
    
    def calculateSumByDate(self,date, act):
        """
        """
        result = {}
        for d, a in zip(date, act):
            if d in result:
                result[d] += a
            else:
                result[d] = a
        return result
    
    def createGraph(self):
        """
        """
        filename = [self.graph_filename1, self.graph_filename2, self.graph_filename3, self.graph_filename4 ]
        return filename

    def handleMessageBoxButton(self,button, file_path):
        if button.text() == "Ok":
            QDesktopServices.openUrl(QUrl.fromLocalFile(file_path))

    def exportWordDoc_2(self):
        try:
            document = Document()

            # Set font
            style = document.styles['Normal']
            style.font.name = 'Calibri'

            # Add header
            header_section = document.sections[0]
            header = header_section.header
            header2 = header_section.header
            
            now = datetime.now()
            date_time = now.strftime("%Y-%m-%d")
            conn = MySQLdb.Connection(host=self.db.DB_SERVER, user=self.db.DB_USERNAME, password=self.db.DB_PASSWORD,
                                database=self.db.DB_NAME)
            mycursor = conn.cursor()
            query = f"SELECT name FROM managers WHERE id = {self.LOGGED_USER_ID}"
            mycursor.execute(query)
            data = mycursor.fetchone()
            name = data[0]
            header_text = header.paragraphs[0]
            header_text.text = str(date_time)
            header_text.alignment = WD_ALIGN_PARAGRAPH.LEFT

            right_header = header2.paragraphs[0]
            right_header.text = str(date_time) + '  ' + name
            right_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            # Add title
            document.add_heading('Tendency of GCINO-EM Project', 0)

            # Add paragraphs
            p = document.add_paragraph('This document presents ')
            p.add_run('the tendency').bold = True
            p.add_run(' of our GCINO application, displaying costs according to projects, users and objects, and finally a final graph of Quantities of objects used according to time.')
            #p.add_run('italic.').italic = True

            titles = ['Amounts by projects','Amounts by users','Amounts by objects','Quantity by date']
            parag = ['This graph shows amounts by project', 'This graph shows amounts by users', 'This graph shows amounts by objects', 'This graph shows quantity by date']
            for i in range(len(titles)):
                graph_filename = self.createGraph()[i]
                document.add_heading(titles[i], level=1)
                document.add_paragraph(parag[i], style='Intense Quote')
                document.add_picture(graph_filename)
                #document.add_picture(graph_filename2)
                if i < len(titles) - 1:
                    document.add_page_break()
            now = datetime.now()
            date_time = now.strftime("%Y_%m_%d")
            conn = MySQLdb.Connection(host=self.db.DB_SERVER, user=self.db.DB_USERNAME, password=self.db.DB_PASSWORD,
                                database=self.db.DB_NAME)
            mycursor = conn.cursor()
            query = f"SELECT name FROM managers WHERE id = {self.LOGGED_USER_ID}"
            mycursor.execute(query)
            data = mycursor.fetchone()
            name = data[0]

            file_dialog = QFileDialog()
            download_path, _ = file_dialog.getSaveFileName(None, "Save Document", "", "Word Document (*.docx)")

            if download_path:
                document.save(download_path)
        except Exception as e:
            QMessageBox.about(self, "Error", "error : "+ str(e))       

class modifyform(insertform, QDialog):
    def __init__(self,table_widget):
        super().__init__()
        self.table_widget = table_widget
        self.table_widget.cellClicked.connect(self.fillForm)
        self.uim = self.ui
        self.uim.actionBox.setEnabled(False)
        self.uim.insertButton.setVisible(False)
        self.uim.updateButton.setVisible(True)
        self.uim.updateButton.clicked.connect(self.updateRow)
        self.uim.cancelButton.clicked.connect(self.cancel)
        self.items = []

    def fillForm(self, row, col):
        """
        """
        self.items.clear()
        for col in range(self.table_widget.columnCount()):
            item = self.table_widget.item(row, col)
            self.items.append(item)

        self.uim.objectText.setText(self.items[0].text())
        self.uim.TypeBox.setCurrentText(self.items[1].text())
        self.uim.locationText.setText(self.items[2].text())
        self.uim.calibrationText.setText(self.items[3].text())
        self.uim.quantityText.setText(self.items[4].text())

    def updateRow(self):
        """
        """
        """
        This function will update a row of data on the object_dist table
        :return:
        """
        objectText      = str(self.uim.objectText.text())
        typeText        = str(self.uim.TypeBox.currentText())
        locationText    = str(self.uim.locationText.text())
        calibrationText = str(self.uim.calibrationText.text())
        quantityText    = str(self.uim.quantityText.text())
        projectText     = str(self.uim.projectBox.currentText())
        categoryText    = str(self.uim.categoryBox.currentText())
        #coutText        = str(self.uim.coutText.text())
        
        old_object = self.items[0].text()
        old_type = self.items[1].text()
        old_location = self.items[2].text()
        old_calibration = self.items[3].text()
        old_quantity = self.items[4].text()
        old_category = self.items[5].text()

        if(len(objectText.strip()) == 0 or len(calibrationText.strip()) == 0 or len(quantityText.strip()) == 0 or len(locationText.strip()) == 0 or len(typeText.strip()) == 0
            or len(projectText.strip()) == 0):
            msgBox = QMessageBox()
            msgBox.setText("All fields must be filled before confirming !")
            msgBox.setWindowTitle("Attention")
            msgBox.exec()
        elif (objectText == old_object) and (typeText == old_type) and (locationText == old_location) and (calibrationText == old_calibration) and (quantityText == old_quantity) and (categoryText == old_category) :
            QMessageBox.about(self, 'Warning', 'This item exists !')
        else:
            conn = MySQLdb.Connection(host=self.db.DB_SERVER, user=self.db.DB_USERNAME, password=self.db.DB_PASSWORD,
                                database=self.db.DB_NAME)
            mycursor = conn.cursor()
            if int(quantityText) == 0:
                query_delete ="DELETE FROM object_dist where Object= '%s' AND Type_object='%s' AND Location='%s' AND Calibration=%s AND Category='%s'" % (objectText,typeText, locationText, int(calibrationText), categoryText)
                mycursor.execute(query_delete)
                conn.commit()
                self.close()
            elif int(quantityText) < 0:
                QMessageBox.about(self, 'Error', 'error : Not valid quantity')
            else: 
                query_update = "UPDATE object_dist SET Object= '%s', Type_object='%s', Location='%s', Calibration='%s', Quantity='%s', Category='%s' WHERE Object = '%s' AND Type_object='%s' AND Location='%s' AND Calibration='%s' AND Quantity='%s' AND Category='%s'" % (objectText,typeText, locationText, calibrationText, quantityText, categoryText, old_object, old_type, old_location, old_calibration, old_quantity, old_category)
                mycursor.execute(query_update)
                conn.commit()
                self.close()

                now = datetime.now()
                date_time = now.strftime("%Y-%m-%d %H:%M:%S")
                query_insert = "INSERT INTO historique(Object,Type_object,Location, Calibration, Quantity, Category,operation_datetime,user_id,actions,project_name) VALUES ('%s','%s','%s','%s','%s','%s','%s','%s','%s','%s')" %(objectText,typeText,locationText,calibrationText,quantityText,categoryText, date_time, self.LOGGED_USER_ID, "Changing", projectText)
                mycursor.execute(query_insert)
                conn.commit()
                self.close()

class bordoreauform(QDialog):
    def __init__(self, table_warehouseman, row):
        """
        """
        super(bordoreauform, self).__init__()
        self.ui = loadUi(os.path.join(os.path.dirname(__file__), "approuved.ui"), self)
        self.table_warehouseman = table_warehouseman
        self.row = row

        #Connect to DB
        self.db = Database()
        config_object = ConfigParser()
        config_object.read("config.ini")
        userInfo = config_object["USERINFO"]
        self.LOGGED_USER_ID = userInfo["LOGGED_USER_ID"]
        
        #Local events
        self.ui.gener_br_btn.clicked.connect(self.addToBordoreau)

    def getNumberOfBordoreau(self):
        """
        """
        try:
            conn = MySQLdb.Connection(host=self.db.DB_SERVER, user=self.db.DB_USERNAME, password=self.db.DB_PASSWORD,
                                    database=self.db.DB_NAME)
            mycursor = conn.cursor()
            
            query_select = """SELECT count(Distinct(user_id)) from orders where status='Approuved'"""
            mycursor.execute(query_select)
            data = mycursor.fetchall()
            for i in data:
                numb_bord = i[0]      
        except Exception as e:
            QMessageBox.about(self, "Warning", str(e))

        return numb_bord
    
    def addToBordoreau(self):
        """
        """
        #self.table_warehouseman.setItem(row,4,QTableWidgetItem("Confirmed"))
        id = self.table_warehouseman.item(self.row, 0).text()
        price = str(self.ui.cout.text())
        source = str(self.ui.link.text())
        quantity = str(self.ui.quantity.text())
        description = str(self.ui.description.toPlainText())

        if not price or not source:
            QMessageBox.about(self, "Warning","Both fields must be filled.")
            return
        #generate or update the excel file
        excel_file = "./test.xlsx"

        workbook = load_workbook(excel_file)
        sheet = workbook.active
        for row_number, row in enumerate(sheet.iter_rows(min_row=5, values_only=True), start=5):
            id_value = row[0]  # Assuming id is in the first column
            if id_value == id:
                # Update the value in the third column (column index is 3)
                cell_price = sheet.cell(row=row_number, column=4)
                cell_price.value = price
                cell_quantity = sheet.cell(row=row_number, column=3)
                cell_quantity.value = quantity
                cell_source = sheet.cell(row=row_number, column=6)
                cell_source.value = source
                cell_description = sheet.cell(row=row_number, column=5)
                cell_description.value = description
        workbook.save(excel_file)

        # Remove the cell widget from the table
        empty_widget = QWidget()
        self.table_warehouseman.setCellWidget(self.row, 6, empty_widget)
        self.table_warehouseman.setCellWidget(self.row, 5, empty_widget)
        self.table_warehouseman.setItem(self.row,4,QTableWidgetItem("Confirmed"))

        try:
            conn = MySQLdb.Connection(host=self.db.DB_SERVER, user=self.db.DB_USERNAME, password=self.db.DB_PASSWORD,
                                    database=self.db.DB_NAME)
            mycursor = conn.cursor()
            
            query = """UPDATE orders SET status='Confirmed', price=%s, quantity=%s WHERE orderID='%s'"""%(float(price), int(quantity), id)
            mycursor.execute(query)
            conn.commit()
            
        except Exception as e:
            QMessageBox.about(self, "Warning", str(e))
        self.close()
